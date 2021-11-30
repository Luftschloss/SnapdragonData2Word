from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches

import json
import os
import re
import sys

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np

MultiGroupConfig = "MultiGroupConfig.json"
MatrixInfoConfig = "SDConfig.json"
XAxisTitle = "Time / s"
MatrixCurveWidth = 0.5
MatrixCurveColor = [[50/255, 177/255, 250/255], [180/255, 144/255, 245/255], [249/255, 196/255, 15/255],
                    [131/255, 211/255, 56/255], [243/255, 81/255, 69/255]]


# 处理数据组，生成Word内容
def processDataCurveByConfig(document: Document, matrixKeyList, matrixDataDic, slicedPeriods):
    # 读取GroupConfig
    groupConfigDic = {}
    with open(MultiGroupConfig, 'r', encoding='utf-8') as f:
        groupConfigDic = json.loads(f.read())
    for group in groupConfigDic.values():
        for matrix in group:
            if not matrixKeyList.__contains__(matrix):
                print("Remove Matrix : " + matrix)
                group.remove(matrix)

    for key, v in groupConfigDic.items():
        document.add_heading(key, level=1)

        # 数据绘图，生成图片
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        plt.rcParams['figure.max_open_warning'] = 40  # 画图最大限制40
        matrixDataArray = []
        for keyName in v:
            for matrixData in matrixDataDic.values():
                if matrixData.keyName == keyName:
                    matrixDataArray.append(matrixData)
                    break
        if len(matrixDataArray) == 0:
            print("Group " + key + " No Data")
            continue
        curveImgPath = drawMultiAndSaveFigure(matrixDataArray, key, slicedPeriods)
        document.add_picture(curveImgPath, width=Inches(6))
        # 生成数据表格
        dataTable = document.add_table(len(matrixDataArray) + 1, 4, style="Light Grid")
        dataTable.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
        headLine = ["检测项", "最小值", "最大值", "平均值"]
        for i in range(4):
            dataTable.cell(0, i).text = headLine[i]
        for i in range(len(matrixDataArray)):
            matrixData = matrixDataArray[i]
            dataLine = [matrixData.matrixNameCN, str(round(matrixData.minValue, 2)), str(round(matrixData.maxValue, 2)),
                        str(round(matrixData.getAverageValue(), 2))]
            for j in range(4):
                dataTable.cell(i + 1, j).text = dataLine[j]


def processDataCurveOneByOne(document: Document, processTypeList, matrixDataDic):
    for i in range(len(processTypeList)):
        document.add_heading(processTypeList[i], level=1)
        for matrixData in matrixDataDic.values():
            if matrixData.processType == processTypeList[i]:
                print(str(len(matrixData.frameDataList)) + '-->' + matrixData.matrixName)
                # 数据绘图，生成图片
                plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
                plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
                plt.rcParams['figure.max_open_warning'] = 40  # 画图最大限制40
                curveImgPath = drawSingleAndSaveFigure(matrixData)
                document.add_heading(matrixData.matrixName, level=2)
                document.add_picture(curveImgPath, width=Inches(6))
                document.add_paragraph('最小值：' + str(round(matrixData.minValue, 2)))
                document.add_paragraph('最大值：' + str(round(matrixData.maxValue, 2)))
                document.add_paragraph('平均值：' + str(round(matrixData.getAverageValue(), 2)))


def createSnapDragonDataDocx(csvDataPath, word_path, timePeriods, slicedPeriods):
    document = Document()
    document.add_heading('Snapdragon 数据', level=0)

    print("Process CSV Begin")
    processTypeList, matrixDataDic, matrixKeyList = csv_process(csvDataPath, timePeriods)
    print("Process CSV End")

    # processDataCurveOneByOne(document, processTypeList, matrixDataDic)
    processDataCurveByConfig(document, matrixKeyList, matrixDataDic, slicedPeriods)

    document.save(word_path)
    return os.getcwd() + '\\' + word_path
    print("Save SnapDragon Docx")


# 每一帧的数据类
class KeyFrameData:
    def __init__(self):
        self.time = 0  # 时间
        self.value = 0  # 数值


# 单个Realtime矩阵数据类
class DataMatrix:
    def __init__(self):
        self.processType = ''  # 进程名称
        self.keyName = ''   #用作Key的唯一名称
        self.matrixName = ''  # 数据名称
        self.matrixNameCN = ''  # 中文名
        self.matrixFileName = ''  # 数据文件名替换了'/'
        self.start = 0  # 数据起始Index
        self.end = 0  # 数据末尾Index
        self.frameDataList = []  # 数据列表
        self.unitStr = 'unit'  # 单位
        self.matrixDescENG = 'desc'  # 英文描述
        self.matrixDescCN = 'desc'  # 中文描述
        self.minValue = sys.float_info.max  # 极小值
        self.maxValue = sys.float_info.min  # 极大值

    def append(self, frameData: KeyFrameData):
        self.frameDataList.append(frameData)
        if frameData.value < self.minValue:
            self.minValue = frameData.value
        if frameData.value > self.maxValue:
            self.maxValue = frameData.value

    def getAverageValue(self):
        valueSum = 0
        for frameData in self.frameDataList:
            valueSum += frameData.value
        return valueSum / len(self.frameDataList)

    def getMaxValue(self):
        return self.maxValue

    def getMinValue(self):
        return self.minValue

    def getTexturePath(self, foldPath: str):
        return foldPath + "/" + self.matrixFileName + ".png"


# 绘制一条曲线并生成曲线图
def drawSingleAndSaveFigure(matrixData: DataMatrix):
    xa = []
    ya = []
    for frameData in matrixData.frameDataList:
        xa.append(frameData.time)
        ya.append(frameData.value)
    fig = plt.figure()
    ax1 = fig.add_subplot(1, 1, 1)
    ax1.set_xlabel(XAxisTitle)  # x轴标签
    ax1.set_ylabel(matrixData.matrixName)  # y轴标签
    # ax1.set_title(matrixData.matrixDescENG)  # 图标标题
    # ax1.text(6, 37, 'test')  # 文本，(6,37)设置文本注释在图片中的坐标
    ax1.grid(linestyle='--', linewidth=1)  # 背景网格
    ax1.plot(xa, ya, color='g', linestyle="-", linewidth=1, label=matrixData.matrixName)
    plt.legend()

    # Save Figure
    figureFoldPath = "Figures"
    folder = os.path.exists(figureFoldPath)
    if not folder:
        os.makedirs(figureFoldPath)
    figureImgPath = matrixData.getTexturePath(figureFoldPath)
    plt.savefig(figureImgPath)
    return figureImgPath


# 绘制多条曲线并生成曲线图
def drawMultiAndSaveFigure(matrixDataArray, groupName, slicedPeriods):
    fig = plt.figure()
    ax1 = fig.add_subplot(1, 1, 1)
    index = 0
    curveLineWidth = MatrixCurveWidth
    maxY = -1
    minY = 100000
    ax1.set_ylabel(groupName)  # y轴标签
    for matrixData in matrixDataArray:
        xa = []
        ya = []
        for frameData in matrixData.frameDataList:
            xa.append(frameData.time)
            ya.append(frameData.value)
            if frameData.value > maxY:
                maxY = frameData.value
            if frameData.value < minY:
                minY = frameData.value
        ax1.set_xlabel(XAxisTitle)  # x轴标
        # ax1.set_title(matrixData.matrixDescENG)  # 图标标题
        # ax1.text(6, 37, 'test')  # 文本，(6,37)设置文本注释在图片中的坐标
        # ax1.grid(linestyle='--', linewidth=1)  # 背景网格
        ax1.plot(xa, ya, color=MatrixCurveColor[index], linestyle="-", linewidth=curveLineWidth,
                 label=matrixData.matrixNameCN)
        index += 1
    plt.legend()
    plt.axis("tight")

    if slicedPeriods is not None:  # add sliced line
        deltaY = (maxY - minY) * 0.02
        for time in slicedPeriods:
            plt.vlines(time, minY - deltaY, maxY + deltaY, colors="r", linestyles="dashed")

    # Save Figure
    figureFoldPath = "Figures"
    folder = os.path.exists(figureFoldPath)
    if not folder:
        os.makedirs(figureFoldPath)
    figureImgPath = figureFoldPath + "/" + str.split(groupName,'（')[0] + ".png"
    plt.savefig(figureImgPath)
    return figureImgPath


# 时间段数据过滤
def CheckInTimePeriods(time, timePeriods):
    for timePeriod in timePeriods:
        if timePeriod[0] < time < timePeriod[1]:
            return True
    return False


# 读取CSV数据
def csv_process(csv_path, timePeriods):
    # 读取MatrixConfig
    matrixConfig = {}
    with open(MatrixInfoConfig, 'r', encoding='utf-8') as e:
        matrixConfig = json.loads(e.read())

    # csv数据表头
    processColStr = 'Process'
    matrixColStr = 'Metric'
    timeStampColStr = "Timestamp"
    timeStampRawColStr = "TimestampRaw"
    valueColStr = 'Value'

    data = pd.read_csv(csv_path, usecols=[processColStr, matrixColStr, timeStampColStr, valueColStr])
    dataLen = data[processColStr].count()

    processList = data[processColStr].tolist()
    matrixList = data[matrixColStr].tolist()
    timeStampList = data[timeStampColStr].tolist()
    valueList = data[valueColStr].tolist()

    matrixDataDic = {}
    processTypeList = []
    matrixKeyList = []
    matrixInfo = None

    for i in range(dataLen):
        key = processList[i] + '_' + matrixList[i]
        if not processTypeList.__contains__(processList[i]):
            processTypeList.append(processList[i])
        time = round(timeStampList[i] / 1000000, 1)

        if timePeriods is None or CheckInTimePeriods(time, timePeriods):
            frameData = KeyFrameData()
            frameData.time = time

            if key not in matrixDataDic:
                data = DataMatrix()
                matrixName = re.sub(u"\\[.*?\\]", "", matrixList[i])
                keyName = matrixName
                if processList[i] == "Global":
                    keyName = "Global_" + keyName
                if not matrixConfig.__contains__(matrixName):
                    continue
                matrixInfo = matrixConfig[matrixName]
                data.matrixNameCN = matrixInfo[0]
                data.keyName = keyName
                data.matrixName = matrixName

                data.matrixFileName = matrixName.replace('/', 'Per')
                data.processType = processList[i]
                matrixDataDic[key] = data
                if not matrixKeyList.__contains__(keyName):
                    matrixKeyList.append(keyName)
            frameData.value = valueList[i] / matrixInfo[3]
            matrixDataDic[key].append(frameData)
    return processTypeList, matrixDataDic, matrixKeyList
