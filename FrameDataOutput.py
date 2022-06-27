import csv
import operator
import os
from typing import List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from enum import Enum
from PIL import Image
from docx.shared import Cm, Pt, RGBColor


class TableTool:
    def make_table_column(tb, cols):
        if len(cols) != len(tb.rows[0].cells):
            return
        tb.autofit = False
        for row_num in range(len(tb.rows)):
            for col_num in range(len(tb.rows[0].cells)):
                tb.cell(row_num, col_num).width = cols[col_num]

    def add_title(tb, x, y, content):
        run = tb.cell(x, y).paragraphs[0].add_run(content)
        run.bold = True
        run.font.color.rgb = RGBColor(79, 129, 189)
        tb.cell(x, y).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


# 每一帧的数据类
class DrawCallData:
    def __init__(self):
        self.ID = 0  # 时间
        self.IDStr = "ID"
        self.Name = 0  # 数值
        self.NameStr = "Name"
        self.Parameters = ""  # 参数
        self.ParametersStr = "Parameters"
        self.Clocks = 0
        self.ClocksStr = "Clocks"
        self.SPRead = 0
        self.SPReadStr = "SP Memory Read (Bytes)"
        self.VertexRead = 0
        self.VertexReadStr = "Vertex Memory Read (Bytes)"
        self.TextureRead = 0
        self.TextureReadStr = "Texture Memory Read BW (Bytes)"
        self.ReadTotal = 0
        self.ReadTotalStr = "Read Total (Bytes)"
        self.WriteTotal = 0
        self.WriteTotalStr = "Write Total (Bytes)"


class ImageType(Enum):
    T2D = 1         #TEXTURE2D
    MIP = 2         #MIPMAP
    CUBE = 4        #CUBEMAP
    T2DA = 8        #TEXTURE2DARRAY
    T3D = 16        #TEXTURE3D


class ImageInfo:
    def __init__(self):
        self.imagePath = ""
        self.imageName = ""
        self.imageType = ImageType.T2D
        self.size = (-1, -1)
        self.isFrameCapture = False


def getAllDrawCalls(csv_path):
    csv_file = open(csv_path)
    csv_reader_lines = csv.reader(csv_file)
    drawCallDataArray = []
    idx = -1
    headline = []
    headIdxDic = {}
    for one_line in csv_reader_lines:
        idx = idx + 1
        if (one_line[0].strip() == '') or (one_line[5].strip() == ''):
            continue
        if idx == 0:
            headline = one_line.copy()
            for i, val in enumerate(headline):
                if val.strip() != "":
                    headIdxDic[val.strip()] = i
            continue
        drawCall = DrawCallData()
        drawCall.ID = int(one_line[headIdxDic[drawCall.IDStr]].strip())
        drawCall.Name = one_line[headIdxDic[drawCall.NameStr]].strip()
        drawCall.Parameters = one_line[headIdxDic[drawCall.ParametersStr]].strip()
        drawCall.Clocks = int(one_line[headIdxDic[drawCall.ClocksStr]].strip())
        drawCall.SPRead = int(one_line[headIdxDic[drawCall.SPReadStr]].strip())/2014
        drawCall.VertexRead = int(one_line[headIdxDic[drawCall.VertexReadStr]].strip())/1024
        drawCall.TextureRead = int(one_line[headIdxDic[drawCall.TextureReadStr]].strip())/1024
        drawCall.WriteTotal = int(one_line[headIdxDic[drawCall.WriteTotalStr]].strip())/1024
        drawCall.ReadTotal = int(one_line[headIdxDic[drawCall.ReadTotalStr]].strip())/1024
        drawCallDataArray.append(drawCall)
    return drawCallDataArray


def getImageInfo(imageFoldPath, fileName: str):
    fullPath = imageFoldPath + "\\" + fileName
    imageInfo = ImageInfo()
    if os.path.isdir(fullPath):
        files = os.listdir(fullPath)
        imageInfo.imagePath = fullPath + "\\" + files[0]
        if fileName.__contains__("(Mipmap)"):
            imageInfo.imageType = ImageType.MIP
            imageInfo.imageName = fileName.strip("(Mipmap)")
        elif fileName.__contains__("(Cubemap)"):
            imageInfo.imageType = ImageType.CUBE
            imageInfo.imageName = fileName.strip("(Cubemap)")
        elif fileName.__contains__("(Texture3D)"):
            imageInfo.imageType = ImageType.T3D
            imageInfo.imageName = fileName.strip("(Texture3D)")
        elif fileName.__contains__("(Texture2DArray)"):
            imageInfo.imageType = ImageType.T2DA
            imageInfo.imageName = fileName.strip("(Texture2DArray)")
    else:
        if fileName.endswith(".png"):
            imageInfo.imagePath = fullPath
            imageInfo.imageName = fileName.strip(".png")
            imageInfo.imageType = ImageType.T2D
            if fileName.startswith("DrawCall_") or fileName.startswith("HighLight_"):
                imageInfo.isFrameCapture = True
        else:
            return None
    if imageInfo.imagePath != "":
        img = Image.open(imageInfo.imagePath)
        imageInfo.size = img.size
        return imageInfo
    else:
        return None


def getDrawCallImages(drawCall, imagePath):
    imageInfoList: List[ImageInfo] = []
    dcImageFoldPath = imagePath + "\\" + str(drawCall)
    if os.path.isdir(dcImageFoldPath):
        files = os.listdir(dcImageFoldPath)
        for file in files:
            imageInfo = getImageInfo(dcImageFoldPath, file)
            if imageInfo is not None:
                imageInfoList.append(imageInfo)
    else:
        print("No Image DrawCall " + str(drawCall))
    return imageInfoList

def getTaleValueStr(value, valueSum):
    return str(round(value, 2)) + " (" + str("%.1f%%" % (value * 100.0 / valueSum)) + ")"

def getTopDrawCall(csv_path, word_path, topNum, Matrix, frameResPath):
    document = Document()
    document.add_heading('Snapdragon FrameData', level=0)
    highLightIdx = -1
    allDrawCalls = getAllDrawCalls(csv_path)
    camp = ()
    if Matrix == "Read Total (Bytes)":
        camp = operator.attrgetter('ReadTotal')
    elif Matrix == "Texture Memory Read BW (Bytes)":
        camp = operator.attrgetter('WriteTotal')
    elif Matrix == "Clocks":
        camp = operator.attrgetter('Clocks')
    elif Matrix == "":
        camp = operator.attrgetter('')

    allDrawCalls.sort(key=camp, reverse=True)

    drawCallSum = DrawCallData()
    topDatasum = 0
    topIdx = 0
    drawCallSum.ID = "Sum"
    for dc in allDrawCalls:
        if topIdx < topNum:
            if Matrix == "Read Total (Bytes)":
                topDatasum += dc.ReadTotal
            elif Matrix == "Texture Memory Read BW (Bytes)":
                topDatasum += dc.WriteTotal
            elif Matrix == "Clocks":
                topDatasum += dc.Clocks
            topIdx += 1
        drawCallSum.Clocks += dc.Clocks
        drawCallSum.VertexRead += dc.VertexRead
        drawCallSum.TextureRead += dc.TextureRead
        drawCallSum.WriteTotal += dc.WriteTotal
        drawCallSum.ReadTotal += dc.ReadTotal

    # 生成Top表格
    document.add_heading('Summary', level=1)
    targetSum = 1
    if Matrix == "Read Total (Bytes)":
        str2 = "的DrawCall带宽（Read）占比为"
        targetSum = drawCallSum.ReadTotal
    elif Matrix == "Texture Memory Read BW (Bytes)":
        str2 = "的DrawCall带宽（Write）占比为"
        targetSum = drawCallSum.WriteTotal
    elif Matrix == "Clocks":
        str2 = "的DrawCall时钟周期（Clocks）占比为"
        targetSum = drawCallSum.Clocks
    summaryParagraph = "单帧按" + Matrix + "排序，Top" + str(topNum) + str2 + str("%.1f%%" % (topDatasum * 100.0 / targetSum)) + "，数据如下表:"
    document.add_paragraph(summaryParagraph, style='Body Text')
    dataTable = document.add_table(topNum + 2, 7, style="Table Grid")
    dataTable.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    headLine = ["Index", "DrawCall", "Clocks", "Vertex Memory Read（KB）", "Texture Memory Read BW（KB）", "Write Total（KB）", "Read Total（KB）"]
    if Matrix == "Clocks":
        highLightIdx = 2
    elif Matrix == "Read Total (Bytes)":
        highLightIdx = 6

    for i in range(7):
        TableTool.add_title(dataTable, 0, i, headLine[i])
    dataTable.cell(topNum + 1, 0).text = drawCallSum.ID
    dataTable.cell(topNum + 1, 1).text = "-"
    dataTable.cell(topNum + 1, 2).text = str(drawCallSum.Clocks)
    dataTable.cell(topNum + 1, 3).text = str(round(drawCallSum.VertexRead, 2))
    dataTable.cell(topNum + 1, 4).text = str(round(drawCallSum.TextureRead, 2))
    dataTable.cell(topNum + 1, 5).text = str(round(drawCallSum.WriteTotal, 2))
    dataTable.cell(topNum + 1, 6).text = str(round(drawCallSum.ReadTotal, 2))
    shading_elm1 = parse_xml(r'<w:shd {} w:fill="FFDE3B"/>'.format(nsdecls('w')))
    dataTable.cell(topNum + 1, highLightIdx)._tc.get_or_add_tcPr().append(shading_elm1)


    for i in range(topNum):
        drawCall = allDrawCalls[i]
        dataLine = [str(i), str(drawCall.ID),
                    getTaleValueStr(drawCall.Clocks, drawCallSum.Clocks),
                    getTaleValueStr(drawCall.VertexRead, drawCallSum.VertexRead),
                    getTaleValueStr(drawCall.TextureRead, drawCallSum.TextureRead),
                    getTaleValueStr(drawCall.WriteTotal, drawCallSum.WriteTotal),
                    getTaleValueStr(drawCall.ReadTotal, drawCallSum.ReadTotal)]
        for j in range(7):
            dataTable.cell(i + 1, j).text = dataLine[j]
        shading_elm2 = parse_xml(r'<w:shd {} w:fill="FFDE3B"/>'.format(nsdecls('w')))
        dataTable.cell(i + 1, highLightIdx)._tc.get_or_add_tcPr().append(shading_elm2)

    retStr = "Top {0} {1} 占比 {2}".format(topNum, str2, str("%.1f%%" % (topDatasum * 100.0 / targetSum)))
    print(retStr)

    # 单个DrawCall数据处理
    fontSize = Pt(10)
    for i in range(topNum):
        dc = allDrawCalls[i]
        document.add_heading("Top " + str(i+1), level=2)
        document.add_paragraph("DrawCall " + str(dc.ID))
        imageInfoList = getDrawCallImages(dc.ID, frameResPath)
        dcTable = document.add_table(2, 3, style="Table Grid")

        dcTable.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
        TableTool.add_title(dcTable, 0, 0, "DrawCall及说明")
        # TableTool.add_title(dcTable, 0, 0, "帧截图")
        TableTool.add_title(dcTable, 0, 1, "渲染相关资源")
        TableTool.add_title(dcTable, 0, 2, "GPU数据")

        isFirstImage = True
        for imageInfo in imageInfoList:
            if imageInfo.isFrameCapture:
                if(imageInfo.imageName.startswith("DrawCall_")):
                    pr1 = dcTable.cell(1,0).paragraphs[0].add_run()
                    pic = pr1.add_picture(imageInfo.imagePath)
                    scale1 = imageInfo.size[0] / 8.8
                    pic.height = Cm(imageInfo.size[1]/scale1)
                    pic.width = Cm(8.8)
                    dcTable.cell(1, 0).add_paragraph("{0}x{1}".format(imageInfo.size[0], imageInfo.size[1]))
                elif(imageInfo.imageName.startswith(("HighLight"))):
                    dcTable.cell(1, 0).add_paragraph()
                    pr1 = dcTable.cell(1, 0).paragraphs[2].add_run()
                    pic = pr1.add_picture(imageInfo.imagePath)
                    scale1 = imageInfo.size[0] / 8.8
                    pic.height = Cm(imageInfo.size[1] / scale1)
                    pic.width = Cm(8.8)
            else:
                p1 = None
                if isFirstImage:
                    isFirstImage = False
                    p1 = dcTable.cell(1,1).paragraphs[0]
                else:
                    p1 = dcTable.cell(1,1).add_paragraph()
                pic = p1.add_run().add_picture(imageInfo.imagePath)
                pic.height = Cm(2)
                pic.width = Cm(2)
                p1.add_run("\n{0}({1})\n{2}x{3}\n".format(imageInfo.imageName, imageInfo.imageType.name, imageInfo.size[0], imageInfo.size[1])).font.size = fontSize

        drawCall = allDrawCalls[i]
        text = dcTable.cell(1, 2).paragraphs[0].add_run("Read:\n"+getTaleValueStr(drawCall.ReadTotal, drawCallSum.ReadTotal) +
                                    "\nWrite:\n" + getTaleValueStr(drawCall.WriteTotal, drawCallSum.WriteTotal) +
                                    "\nClocks:\n" + getTaleValueStr(drawCall.Clocks, drawCallSum.Clocks) +
                                    "\nTexture:\n" + getTaleValueStr(drawCall.TextureRead, drawCallSum.TextureRead) +
                                    "\nVertex:\n" + getTaleValueStr(drawCall.VertexRead, drawCallSum.VertexRead))
        text.font.size = fontSize

        #Shader Info
        drawCallDataPath = frameResPath + "\\" + str(drawCall.ID) + "\\DrawCallData.txt"
        if os.path.exists(drawCallDataPath):
            f = open(drawCallDataPath)
            shaderInfoStr = f.readlines()[0]
            shaderInfoStr = shaderInfoStr.strip('ID').strip()
            shaderStart = shaderInfoStr.find('/') + 1
            shaderEnd = shaderInfoStr.find(', ')
            shaderName = shaderInfoStr[shaderStart:shaderEnd-1]
            programProperties = shaderInfoStr[shaderEnd+2:].split(', ')
            p = dcTable.cell(1, 1).add_paragraph()
            p.add_run("Shader：{0}\n".format(shaderName)).font.size = fontSize
            p.add_run("\nProperties：\n").font.size = fontSize
            for prop in programProperties:
                p.add_run("{0}\n".format(prop)).font.size = fontSize

        widths = (Cm(9.2), Cm(3), Cm(3.2))
        TableTool.make_table_column(dcTable, widths)
    document.save(word_path)

    print("Save SDFrameData")
    return os.getcwd() + '\\' + word_path, retStr
